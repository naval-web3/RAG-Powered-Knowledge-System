"""
Build the MCSP-232 project report.

Reads the Markdown chapters in report/src/ (in filename order) and writes one
Word document formatted to the IGNOU project guidelines: A4, single line
spacing, printed on one side, with every page, figure and table numbered and
every figure and table titled.

    ..\\backend\\.venv\\Scripts\\python.exe build.py           # build report.docx
    ..\\backend\\.venv\\Scripts\\python.exe build.py --stats   # counts only

The Markdown accepted here is a deliberately small subset. Everything the
report actually needs is supported; nothing else is.

    # Title                chapter heading   -> "Chapter 3  System Design"
    ## Title               section           -> "3.2  Modularisation"
    ### Title              sub-section       -> "3.2.1  The RAG engine"
    #### Title             run-in heading, unnumbered
    # *Title               heading with no number (front matter, appendices)

    Plain paragraphs, "- " bullets, "1. " numbered lists, "> " notes.
    **bold**, *italic*, `code` inline.

    ![Caption](assets/x.png){width=5.5}  figure, auto-numbered per chapter
    | a | b |                            pipe table, auto-numbered
    Table: Caption                       titles the table that follows it
    ```python ... ```                    code listing in a shaded box

    <!-- toc -->   contents field      <!-- pagebreak -->   page break
    <!-- lof -->   list of figures     <!-- arabic -->      restart at page 1
    <!-- lot -->   list of tables      <!-- landscape -->   wide section
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SRC = HERE / "src"
OUT = HERE / "report.docx"

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
BODY_SIZE = Pt(12)
CODE_SIZE = Pt(8.5)
TEXT_WIDTH_IN = 6.0  # A4 (8.27") less a 1.25" binding margin and 1" outer margin


# ---------------------------------------------------------------------------
# low-level Word helpers
# ---------------------------------------------------------------------------

def _field(paragraph, instruction: str, placeholder: str) -> None:
    """Insert a Word field (TOC, PAGE, ...) that Word fills in when it updates."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, sep, text, end):
        run._r.append(node)
    return run


def _shade(element, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _cell_borders(cell, colour: str = "9A9A9A", size: int = 4) -> None:
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        line = OxmlElement("w:" + edge)
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), str(size))
        line.set(qn("w:color"), colour)
        borders.append(line)
    cell._tc.get_or_add_tcPr().append(borders)


_BOOKMARK_ID = [1000]


def _bookmark(paragraph, name: str) -> None:
    """Wrap a paragraph in a bookmark so a PAGEREF field can point at its page."""
    _BOOKMARK_ID[0] += 1
    ident = str(_BOOKMARK_ID[0])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), ident)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), ident)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _keep_together(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.keep_together = True
    pf.keep_with_next = True


def _number_format(section, fmt: str, start: int | None = None) -> None:
    """Page-number format for a section: 'decimal' or 'lowerRoman'."""
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    section._sectPr.append(pg)


def _page_footer(section, first_page_blank: bool = False) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    _field(para, " PAGE   \\* MERGEFORMAT ", "1")
    for run in para.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(10)


def _running_header(section, text: str) -> None:
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "BFBFBF")
    bottom.set(qn("w:space"), "2")
    borders.append(bottom)
    para._p.get_or_add_pPr().append(borders)


def _setup_page(section) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.25)   # binding edge
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


# ---------------------------------------------------------------------------
# styles
# ---------------------------------------------------------------------------

def build_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    specs = {
        "Heading 1": (16.0, 18, 12),
        "Heading 2": (13.5, 14, 6),
        "Heading 3": (12.0, 10, 4),
        "Heading 4": (12.0, 8, 2),
    }
    for name, (size, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = BODY_FONT
        # Word's built-in heading styles name a *theme* font, and a theme font
        # beats the plain w:ascii the line above sets - which is why headings
        # came out in Calibri while the body was Times. Strip the theme
        # attributes so the font actually asked for is the font used.
        fonts = style.element.rPr.rFonts
        for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            if fonts.get(qn("w:" + attribute)) is not None:
                del fonts.attrib[qn("w:" + attribute)]
        fonts.set(qn("w:ascii"), BODY_FONT)
        fonts.set(qn("w:hAnsi"), BODY_FONT)
        fonts.set(qn("w:cs"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = name == "Heading 4"
        style.font.color.rgb = RGBColor(0, 0, 0)
        hpf = style.paragraph_format
        hpf.space_before = Pt(before)
        hpf.space_after = Pt(after)
        hpf.line_spacing = 1.0
        hpf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hpf.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(10)
    caption.font.bold = False
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    cpf = caption.paragraph_format
    cpf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cpf.space_before = Pt(4)
    cpf.space_after = Pt(14)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE
        lpf = style.paragraph_format
        lpf.line_spacing = 1.0
        lpf.space_after = Pt(3)
        lpf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lpf.left_indent = Inches(0.35)


# ---------------------------------------------------------------------------
# inline markdown
# ---------------------------------------------------------------------------

INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*\n]+?\*(?!\*)|`[^`]+?`)")


def add_inline(paragraph, text: str, size=None, code_size=None):
    """Write text into a paragraph honouring **bold**, *italic* and `code`."""
    size = size or BODY_SIZE
    code_size = code_size or Pt(size.pt - 1.5)
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            run = paragraph.add_run(piece[1:-1])
            run.font.name = CODE_FONT
            run.font.size = code_size
            continue
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            run = paragraph.add_run(piece[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(piece)
        run.font.size = size
    return paragraph


# ---------------------------------------------------------------------------
# the builder
# ---------------------------------------------------------------------------

class ReportBuilder:
    def __init__(self, doc: Document):
        self.doc = doc
        self.chapter = 0
        self.section_no = 0
        self.subsection_no = 0
        self.figure_no = 0
        self.table_no = 0
        self.numbered = True          # front matter switches this off
        self.pending_table_caption = None
        self.figures: list[str] = []
        self.tables: list[str] = []
        self.words = 0
        self.started = False          # no page break before the very first block
        self.markers: dict[str, object] = {}   # where the figure/table lists go

    # -- blocks ------------------------------------------------------------

    def para(self, text: str, style: str | None = None):
        self.started = True
        p = self.doc.add_paragraph(style=style)
        add_inline(p, text)
        self.words += len(text.split())
        return p

    def heading(self, level: int, text: str):
        unnumbered = text.startswith("*")
        if unnumbered:
            text = text[1:].strip()

        if level == 1:
            if self.started:
                self.doc.add_page_break()
            self.started = True
            self.section_no = 0
            self.subsection_no = 0
            title = text.upper()
            if not unnumbered and self.numbered:
                self.chapter += 1
                self.figure_no = 0
                self.table_no = 0
                # The chapter number belongs in the heading text itself, not in
                # a separate line above it: Word builds the contents from
                # heading paragraphs, and anything outside them is invisible to
                # it. "CHAPTER 3  SYSTEM DESIGN" reads correctly in both places.
                title = "CHAPTER %d.  %s" % (self.chapter, title)
            head = self.doc.add_heading(title, level=1)
            self._rule(head)
            return head

        if level == 2:
            self.section_no += 1
            self.subsection_no = 0
            prefix = "%d.%d  " % (self.chapter, self.section_no) if self.numbered and not unnumbered else ""
            return self.doc.add_heading(prefix + text, level=2)

        if level == 3:
            self.subsection_no += 1
            prefix = (
                "%d.%d.%d  " % (self.chapter, self.section_no, self.subsection_no)
                if self.numbered and not unnumbered
                else ""
            )
            return self.doc.add_heading(prefix + text, level=3)

        return self.doc.add_heading(text, level=4)

    def _rule(self, paragraph) -> None:
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "333333")
        bottom.set(qn("w:space"), "4")
        borders.append(bottom)
        paragraph._p.get_or_add_pPr().append(borders)

    def figure(self, path: str, caption: str, width: float | None) -> None:
        image = HERE / path
        if not image.exists():
            print("  !! missing figure: %s" % path, file=sys.stderr)
            return
        self.figure_no += 1
        number = "%d.%d" % (self.chapter, self.figure_no)
        holder = self.doc.add_paragraph()
        holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        holder.paragraph_format.space_before = Pt(10)
        holder.paragraph_format.space_after = Pt(2)
        holder.paragraph_format.keep_with_next = True
        run = holder.add_run()
        run.add_picture(str(image), width=Inches(width or TEXT_WIDTH_IN))
        cap = self.doc.add_paragraph(style="Caption")
        add_inline(cap, "Figure %s: %s" % (number, caption), size=Pt(10))
        mark = "fig_%s" % number.replace(".", "_")
        _bookmark(cap, mark)
        self.figures.append(("Figure %s: %s" % (number, caption), mark))

    def table(self, rows: list[list[str]], caption: str | None) -> None:
        self.table_no += 1
        number = "%d.%d" % (self.chapter, self.table_no)
        title = caption or "Untitled"
        cap = self.doc.add_paragraph(style="Caption")
        cap.paragraph_format.space_before = Pt(10)
        cap.paragraph_format.space_after = Pt(3)
        cap.paragraph_format.keep_with_next = True
        add_inline(cap, "Table %s: %s" % (number, title), size=Pt(10))
        mark = "tbl_%s" % number.replace(".", "_")
        _bookmark(cap, mark)
        self.tables.append(("Table %s: %s" % (number, title), mark))

        table = self.doc.add_table(rows=0, cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for r, source in enumerate(rows):
            cells = table.add_row().cells
            for c, text in enumerate(source):
                if c >= len(cells):
                    break
                cell = cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline(p, text.strip(), size=Pt(10))
                if r == 0:
                    for run in p.runs:
                        run.bold = True
                    _shade(cell._tc.get_or_add_tcPr(), "E8E8E8")
                _cell_borders(cell)
                self.words += len(text.split())
        if table.rows:
            trPr = table.rows[0]._tr.get_or_add_trPr()
            header = OxmlElement("w:tblHeader")
            trPr.append(header)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def code(self, lines: list[str], language: str) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        cell.text = ""
        _shade(cell._tc.get_or_add_tcPr(), "F4F4F4")
        _cell_borders(cell, "C8C8C8", 4)
        first = True
        for line in lines:
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line.replace("\t", "    "))
            run.font.name = CODE_FONT
            run.font.size = CODE_SIZE
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def note(self, text: str) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = ""
        _shade(cell._tc.get_or_add_tcPr(), "F2F4F7")
        _cell_borders(cell, "B9C2CE", 4)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_inline(p, text, size=Pt(11))
        self.words += len(text.split())
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# the one page that is laid out by hand
# ---------------------------------------------------------------------------

TITLE_PAGE = [
    ("A PROJECT REPORT ON", 12, False, 26),
    ("RAG POWERED KNOWLEDGE SYSTEM", 20, True, 30),
    ("Submitted to the School of Computer and Information Sciences, IGNOU "
     "in partial fulfilment of the requirements for the award of the degree", 11, False, 22),
    ("MASTER OF COMPUTER APPLICATIONS", 15, True, 2),
    ("(MCA_NEW)", 12, False, 2),
    ("MCSP-232", 12, True, 30),
    ("SUBMITTED TO", 11, False, 6),
    ("INDIRA GANDHI NATIONAL OPEN UNIVERSITY", 13, True, 2),
    ("MAIDAN GARHI, NEW DELHI - 110068", 11, False, 10),
    ("Study Centre Code: 1105          Regional Centre: 11 - SHIMLA", 11, False, 96),
]


def title_page(builder: ReportBuilder) -> None:
    """The cover. Laid out here rather than in Markdown because it is the one
    page in the report that is centred, spaced by eye and never reflows."""
    doc = builder.doc
    builder.started = True
    for text, size, bold, after in TITLE_PAGE:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(after)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = Pt(size)
        run.bold = bold

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = table.rows[0].cells
    blocks = (
        (left, ["Submitted By", "", "Name: NAVAL CHAUDHARY",
                "Enrolment No.: 2354558202", "Programme: MCA_NEW"]),
        (right, ["Guided By", "", "Name: ____________________",
                 "Signature: ________________", "Date: _____________________"]),
    )
    for cell, lines in blocks:
        cell.text = ""
        for index, line in enumerate(lines):
            para = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(line)
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            run.bold = index == 0


def sign_block(builder: ReportBuilder, arg: str) -> None:
    """A two-column signature strip: '<!-- signblock Left text | Right text -->'
    with lines separated by ' / '."""
    doc = builder.doc
    builder.started = True
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    table = doc.add_table(rows=1, cols=2)
    columns = [side.strip() for side in (arg or "|").split("|")]
    while len(columns) < 2:
        columns.append("")
    for cell, block in zip(table.rows[0].cells, columns):
        cell.text = ""
        lines = [piece.strip() for piece in block.split(" / ")] if block else [""]
        for index, line in enumerate(lines):
            para = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(2)
            add_inline(para, line, size=Pt(11))


def centred(builder: ReportBuilder, arg: str) -> None:
    builder.started = True
    para = builder.doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    add_inline(para, arg or "")


# ---------------------------------------------------------------------------
# markdown parser
# ---------------------------------------------------------------------------

FIGURE = re.compile(r"^!\[(?P<cap>.*?)\]\((?P<path>[^)]+)\)(?:\{width=(?P<w>[\d.]+)\})?\s*$")
HEADING = re.compile(r"^(?P<hashes>#{1,4})\s+(?P<text>.+?)\s*$")
BULLET = re.compile(r"^[-*]\s+(?P<text>.+?)\s*$")
NUMBERED = re.compile(r"^\d+[.)]\s+(?P<text>.+?)\s*$")
TABLE_CAPTION = re.compile(r"^Table:\s*(?P<text>.+?)\s*$")
DIRECTIVE = re.compile(r"^<!--\s*(?P<name>[a-z-]+)(?:\s+(?P<arg>.*?))?\s*-->$")


def parse(builder: ReportBuilder, text: str) -> None:
    doc = builder.doc
    lines = text.replace("\r\n", "\n").split("\n")
    i = 0
    caption = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        directive = DIRECTIVE.match(stripped)
        if directive:
            handle_directive(builder, directive.group("name"), directive.group("arg"))
            i += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            builder.code(block, language)
            continue

        heading = HEADING.match(stripped)
        if heading:
            builder.heading(len(heading.group("hashes")), heading.group("text"))
            i += 1
            continue

        figure = FIGURE.match(stripped)
        if figure:
            width = float(figure.group("w")) if figure.group("w") else None
            builder.figure(figure.group("path"), figure.group("cap"), width)
            i += 1
            continue

        table_caption = TABLE_CAPTION.match(stripped)
        if table_caption:
            caption = table_caption.group("text")
            i += 1
            continue

        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip().strip("|")
                cells = [c.strip() for c in row.split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                    rows.append(cells)
                i += 1
            if rows:
                builder.table(rows, caption)
            caption = None
            continue

        if stripped.startswith("> "):
            block = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                block.append(lines[i].strip()[2:])
                i += 1
            builder.note(" ".join(block))
            continue

        bullet = BULLET.match(stripped)
        if bullet:
            while i < len(lines) and BULLET.match(lines[i].strip()):
                item = BULLET.match(lines[i].strip()).group("text")
                builder.para(item, style="List Bullet")
                i += 1
            continue

        numbered = NUMBERED.match(stripped)
        if numbered:
            while i < len(lines) and NUMBERED.match(lines[i].strip()):
                item = NUMBERED.match(lines[i].strip()).group("text")
                builder.para(item, style="List Number")
                i += 1
            continue

        block = []
        while i < len(lines) and lines[i].strip() and not _starts_block(lines[i].strip()):
            block.append(lines[i].strip())
            i += 1
        builder.para(" ".join(block))

    return None


def _starts_block(stripped: str) -> bool:
    return bool(
        stripped.startswith(("#", "|", "> ", "```", "!["))
        or BULLET.match(stripped)
        or NUMBERED.match(stripped)
        or DIRECTIVE.match(stripped)
        or TABLE_CAPTION.match(stripped)
    )


def handle_directive(builder: ReportBuilder, name: str, arg: str | None) -> None:
    doc = builder.doc
    if name == "pagebreak":
        doc.add_page_break()
    elif name == "frontmatter":
        builder.numbered = False
    elif name == "numbered":
        builder.numbered = True
    elif name == "toc":
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        _field(para, ' TOC \\o "1-3" \\h \\z \\u ', "Right-click and choose Update Field to build the contents.")
    elif name == "plainheading":
        # A heading that looks like a chapter heading but is deliberately not
        # one: the contents, the list of figures and the list of tables should
        # not list themselves.
        builder.started = True
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.keep_with_next = True
        run = para.add_run((arg or "").upper())
        run.font.name = BODY_FONT
        run.font.size = Pt(16)
        run.bold = True
        builder._rule(para)
    elif name in ("lof", "lot"):
        # Figures and tables are numbered by this script, not by Word's SEQ
        # fields, so Word cannot collect them itself. Leave a marker here and
        # fill the list in once every chapter has been parsed and every caption
        # carries a bookmark for its page number to point at.
        builder.markers[name] = doc.add_paragraph()
    elif name == "titlepage":
        title_page(builder)
    elif name == "signblock":
        sign_block(builder, arg or "")
    elif name == "center":
        centred(builder, arg or "")
    elif name == "vspace":
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(float(arg or 12))
    elif name == "arabic":
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        _setup_page(section)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        _number_format(section, "decimal", start=1)
        _page_footer(section)
        _running_header(section, "RAG Powered Knowledge System")
        # The section break already turned the page, so the chapter heading
        # that follows must not turn it again and leave a blank sheet.
        builder.started = False
    elif name == "landscape":
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        _setup_page(section)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Inches(11.69), Inches(8.27)
        _page_footer(section)
    elif name == "portrait":
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        _setup_page(section)
        section.orientation = WD_ORIENT.PORTRAIT
        _page_footer(section)
        _running_header(section, "RAG Powered Knowledge System")
    else:
        print("  !! unknown directive: %s" % name, file=sys.stderr)


def fill_lists(builder: ReportBuilder) -> None:
    """Write the List of Figures and List of Tables into the markers left in
    the front matter, now that every caption exists and carries a bookmark."""
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

    pairs = (("lof", builder.figures), ("lot", builder.tables))
    for key, entries in pairs:
        marker = builder.markers.get(key)
        if marker is None:
            continue
        for label, mark in entries:
            para = builder.doc.add_paragraph()
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Inches(0.0)
            para.paragraph_format.tab_stops.add_tab_stop(
                Inches(TEXT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
            )
            run = para.add_run(label)
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            tab = para.add_run("\t")
            tab.font.size = Pt(11)
            _field(para, " PAGEREF %s \\h " % mark, "0")
            for field_run in para.runs[-1:]:
                field_run.font.name = BODY_FONT
                field_run.font.size = Pt(11)
            marker._p.addprevious(para._p)
        marker._p.getparent().remove(marker._p)


# ---------------------------------------------------------------------------
# entry point
# ---------------------------------------------------------------------------

def main() -> int:
    sources = sorted(SRC.glob("*.md"))
    if not sources:
        print("no chapters in %s" % SRC, file=sys.stderr)
        return 1

    doc = Document()
    build_styles(doc)
    first = doc.sections[0]
    _setup_page(first)
    _number_format(first, "lowerRoman", start=1)
    _page_footer(first)
    # The title page carries no page number, though it counts as page i.
    first.different_first_page_header_footer = True

    builder = ReportBuilder(doc)
    builder.numbered = False

    for path in sources:
        print("  %s" % path.name)
        parse(builder, path.read_text(encoding="utf-8"))

    fill_lists(builder)

    if "--stats" in sys.argv:
        print("\n  words   %d" % builder.words)
        print("  figures %d" % len(builder.figures))
        print("  tables  %d" % len(builder.tables))
        return 0

    doc.save(OUT)
    print("\n  wrote %s" % OUT)
    print("  %d words, %d figures, %d tables, %d chapters"
          % (builder.words, len(builder.figures), len(builder.tables), builder.chapter))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
