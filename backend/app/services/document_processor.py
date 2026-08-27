"""
Document processing pipeline (synopsis Module 3).

Flow: load raw file -> extract text -> chunk (Recursive splitter,
1000 chars / 200 overlap) -> embed -> store in ChromaDB with metadata
-> update the document's processing_status in PostgreSQL.
"""

import uuid

from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy.orm import Session

from app.config import settings
from app.models import Document
from app.services import vector_store


# ---------- OCR (scanned / image-only PDFs) ----------
# RapidOCR is heavy to construct (loads onnx models), so build it once and
# reuse it across documents.
_ocr_engine = None


def _get_ocr_engine():
    global _ocr_engine
    if _ocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR

        _ocr_engine = RapidOCR()
    return _ocr_engine


def _ocr_pdf_pages(path: str, page_indices: list[int], on_progress=None) -> dict[int, str]:
    """OCR the given 0-based page indices of a PDF. Renders each page to an
    image with PyMuPDF and reads it with RapidOCR. Returns {index: text}."""
    import fitz  # PyMuPDF
    import numpy as np
    from PIL import Image

    engine = _get_ocr_engine()
    zoom = settings.OCR_DPI / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    out: dict[int, str] = {}
    doc = fitz.open(path)
    targets = page_indices[: settings.OCR_MAX_PAGES]
    try:
        for n_done, idx in enumerate(targets, start=1):
            pix = doc[idx].get_pixmap(matrix=matrix, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            result, _ = engine(np.array(img))
            # result is a list of [box, text, score]; keep the recognized text.
            out[idx] = "\n".join(line[1] for line in result) if result else ""
            if on_progress is not None:
                on_progress(n_done, len(targets))
    finally:
        doc.close()
    return out


# ---------- Text extraction ----------
def _extract_pdf(path: str, on_ocr_start=None, on_page=None) -> list[tuple[int, str]]:
    """Return (page_number, text) for each page. Pages with no embedded
    (selectable) text fall back to OCR when OCR_ENABLED. If OCR is triggered,
    on_ocr_start(n_pages) is called first so the caller can flag the slow step.
    on_page(done, total) fires as pages are read, so the caller can report
    real extraction progress rather than guessing at it."""
    from pypdf import PdfReader

    reader = PdfReader(path)
    total = len(reader.pages)
    pages: list[tuple[int, str]] = []
    empty_indices: list[int] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            empty_indices.append(len(pages))
        pages.append((i, text))
        if on_page is not None:
            on_page(i, total)

    # OCR only the pages that yielded no text, so normal (text-based) PDFs and
    # the text pages of mixed PDFs stay fast and are unaffected.
    if settings.OCR_ENABLED and empty_indices:
        if on_ocr_start is not None:
            on_ocr_start(min(len(empty_indices), settings.OCR_MAX_PAGES))
        for idx, ocr_text in _ocr_pdf_pages(path, empty_indices, on_progress=on_page).items():
            pages[idx] = (pages[idx][0], ocr_text)
    return pages


def _extract_docx(path: str) -> list[tuple[int, str]]:
    import docx

    doc = docx.Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return [(1, text)]


def _extract_txt(path: str) -> list[tuple[int, str]]:
    with open(path, encoding="utf-8", errors="ignore") as f:
        return [(1, f.read())]


def extract_text(
    path: str, file_type: str, on_ocr_start=None, on_page=None
) -> list[tuple[int, str]]:
    if file_type == "pdf":
        return _extract_pdf(path, on_ocr_start=on_ocr_start, on_page=on_page)
    if file_type == "docx":
        pages = _extract_docx(path)
    elif file_type == "txt":
        pages = _extract_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    # Single-unit formats have one 'page'; report it so progress still moves.
    if on_page is not None:
        on_page(1, 1)
    return pages


def _clean(text: str) -> str:
    # Light preprocessing: collapse excess whitespace, keep content intact.
    return " ".join(text.split())


# ---------- Section / heading detection ----------
import re

_HEADING_NUM = re.compile(r"^(?:\d+(?:\.\d+)*\.?|[IVXLC]+\.?)\s+\S")


def _is_heading(line: str) -> bool:
    """Best-effort detection of a section heading line."""
    s = line.strip()
    if len(s) < 3 or len(s) > 70:
        return False
    if s[-1] in ".,;:":
        return False
    words = s.split()
    if len(words) > 10:
        return False
    letters = [c for c in s if c.isalpha()]
    if len(letters) < 3:
        return False
    # ALL-CAPS heading, e.g. "ASSESSMENT GUIDELINES FOR PROJECT EVALUATION"
    if all(c.isupper() for c in letters):
        return True
    # Numbered / roman-numeral heading, e.g. "1. Project Proposal" / "VII Assessment"
    if _HEADING_NUM.match(s) and len(words) <= 9:
        return True
    # Short Title Case heading, e.g. "Assessment Criteria"
    if len(words) <= 6 and sum(1 for w in words if w[:1].isupper()) >= max(1, len(words) - 1):
        return True
    return False


def _split_sections(text: str) -> list[tuple[str | None, str]]:
    """Split a page's text into (section_title, body) segments by heading lines."""
    lines = text.split("\n")
    sections: list[tuple[str | None, str]] = []
    current_title: str | None = None
    buf: list[str] = []
    for line in lines:
        if _is_heading(line):
            if buf:
                sections.append((current_title, "\n".join(buf)))
                buf = []
            current_title = line.strip()
        buf.append(line)
    if buf:
        sections.append((current_title, "\n".join(buf)))
    return sections or [(None, text)]


# ---------- Pipeline ----------
# Chunks are embedded and indexed in batches so progress can be reported from
# real work completed rather than estimated from elapsed time.
_EMBED_BATCH = 64

# Fine stage -> coarse processing_status. The coarse value still drives the
# table badge and the status filters, and is the only thing the documents
# status CHECK constraint allows, so it stays limited to those five values.
_COARSE = {
    "extracting": "processing",
    "ocr": "ocr",
    "chunking": "processing",
    "embedding": "processing",
    "indexing": "processing",
    "done": "done",
    "failed": "failed",
}


class _Progress:
    """Records where a document currently is in the ingest pipeline.

    Each stage is given the slice of the 0-100 bar running from wherever the
    previous stage finished up to its own END. Allocating the start dynamically
    (rather than from a fixed table) is what keeps the bar continuous whether or
    not OCR runs: a scanned PDF finishes reading at 35 and OCR then owns 35->60,
    while a text PDF simply lets chunking take that room instead. A fixed table
    would either leave a visible jump or, worse, hand OCR a span the extraction
    pass had already consumed -- which silently froze the bar for the whole of
    the slowest phase.
    """

    # Percent at which each stage finishes.
    END = {"extracting": 35, "ocr": 60, "chunking": 68, "embedding": 85, "indexing": 99}

    def __init__(self, db, document: Document) -> None:
        self.db = db
        self.doc = document
        self.start = 0

    def __call__(self, stage: str, frac: float = 0.0,
                 detail: str | None = None, force: bool = False) -> None:
        """frac is 0..1 *within* `stage`. Progress only ever moves forward, so
        the client can interpolate between polls without the bar going backwards."""
        doc = self.doc
        if stage != doc.stage:
            self.start = doc.progress or 0
        end = self.END.get(stage, 100)
        lo = min(self.start, end)
        pct = int(lo + (end - lo) * max(0.0, min(1.0, frac)))
        pct = max(doc.progress or 0, pct)
        # Nothing visible changed -> skip the write, so a 500-page document does
        # not cause 500 commits.
        if not force and stage == doc.stage and pct == (doc.progress or 0):
            return
        doc.stage = stage
        doc.processing_status = _COARSE.get(stage, "processing")
        doc.progress = pct
        if detail is not None:
            doc.stage_detail = detail[:120]
        self.db.commit()


def process_document(db: Session, document: Document) -> None:
    """Run the full ingest pipeline for one document and persist live progress."""
    # 4, not 0: the file is already accepted and stored by this point, so the
    # bar should not read empty while the first page is being opened.
    document.progress = 4
    document.stage_detail = None
    document.error_message = None
    report = _Progress(db, document)
    report("extracting", 0.0, "reading file", force=True)

    # Set by _mark_ocr so the page callback knows to report OCR rather than
    # plain text extraction for the pages that still have to be read.
    ocr_active = {"on": False}

    def _mark_ocr(n_pages: int) -> None:
        # A scanned / image-only PDF was detected: OCR is about to run, which is
        # slow. Flag it so the UI can tell the user this document is being
        # scanned and will take longer.
        ocr_active["on"] = True
        report("ocr", 0.0,
            f"scanned PDF - running OCR on {n_pages} page(s)", force=True,
        )

    def _on_page(done: int, total: int) -> None:
        stage = "ocr" if ocr_active["on"] else "extracting"
        verb = "OCR" if ocr_active["on"] else "read"
        report(stage, done / max(total, 1),
            f"{verb} {done} of {total} page(s)",
        )

    try:
        pages = extract_text(
            document.file_path,
            document.file_type,
            on_ocr_start=_mark_ocr,
            on_page=_on_page,
        )
        report("chunking", 0.0,
            f"{len(pages)} page(s) extracted", force=True,
        )

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

        chunks: list[LCDocument] = []
        ids: list[str] = []
        for p_no, (page_number, raw) in enumerate(pages, start=1):
            chunk_idx = 0
            for section_title, body in _split_sections(raw):
                cleaned = _clean(body)
                if not cleaned:
                    continue
                for piece in splitter.split_text(cleaned):
                    chunk_id = str(uuid.uuid4())
                    chunks.append(
                        LCDocument(
                            page_content=piece,
                            metadata={
                                "document_id": str(document.document_id),
                                "user_id": str(document.user_id),
                                "title": document.title,
                                "file_type": document.file_type,
                                "page_number": page_number,
                                # Chroma metadata cannot be None -> use "".
                                "section": section_title or "",
                                "chunk_index": chunk_idx,
                            },
                        )
                    )
                    ids.append(chunk_id)
                    chunk_idx += 1
            report("chunking", p_no / max(len(pages), 1),
                f"{len(chunks)} chunks from {p_no} of {len(pages)} page(s)",
            )

        if not chunks:
            if document.file_type == "pdf":
                if settings.OCR_ENABLED:
                    raise ValueError(
                        "No text could be read from this PDF. It has no selectable "
                        "text and OCR couldn't recognize any text in the page images "
                        "— the scan may be blank, very low quality, or rotated. Try a "
                        "clearer scan or a text-based PDF."
                    )
                raise ValueError(
                    "No selectable text found — this looks like a scanned or "
                    "image-only PDF. OCR is disabled, so please upload a "
                    "text-based PDF (one where you can select/copy the text)."
                )
            raise ValueError("No extractable text found in the document.")

        # Embed first, then index, both in batches. These are two distinct
        # pieces of real work, and reporting them separately is what lets the UI
        # show true per-batch progress instead of one long opaque pause.
        total = len(chunks)
        texts = [c.page_content for c in chunks]
        metadatas = [c.metadata for c in chunks]
        bounds = [(i, min(i + _EMBED_BATCH, total)) for i in range(0, total, _EMBED_BATCH)]

        vectors: list[list[float]] = []
        for n, (a, b) in enumerate(bounds, start=1):
            vectors.extend(vector_store.embed_texts(texts[a:b]))
            report("embedding", n / len(bounds),
                f"batch {n} of {len(bounds)} - {b} of {total} chunks embedded",
            )

        for n, (a, b) in enumerate(bounds, start=1):
            vector_store.index_embedded(ids[a:b], texts[a:b], metadatas[a:b], vectors[a:b])
            # Publish the count as it grows so the library shows it climbing.
            document.chunk_count = b
            report("indexing", n / len(bounds),
                f"batch {n} of {len(bounds)} - {b} of {total} chunks indexed",
                force=True,
            )

        document.chunk_count = total
        document.error_message = None
        report("done", 1.0,
            f"{total} chunks ready for retrieval", force=True,
        )

    except Exception as exc:  # noqa: BLE001 - record failure for the user/admin
        document.error_message = str(exc)[:500]
        # Leaves `progress` where it stopped, so the UI can mark the failure on
        # the step that actually broke instead of resetting the whole bar.
        report("failed", 0.0, force=True)
        raise
